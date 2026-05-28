# app.py
import os
import json
import re
from io import BytesIO
from pathlib import Path

# Must be set BEFORE importing crewai / crew
os.environ["CREWAI_DISABLE_TELEMETRY"] = "true"
os.environ["OTEL_SDK_DISABLED"] = "true"

from dotenv import load_dotenv
import streamlit as st

def load_streamlit_secrets():
    try:
        for key, value in st.secrets.items():
            if isinstance(value, (str, int, float, bool)):
                os.environ.setdefault(key, str(value))
    except Exception:
        pass

# Load .env reliably even when Streamlit changes working dir
load_dotenv(dotenv_path=Path(__file__).with_name(".env"))

# Load Streamlit secrets into environment variables
load_streamlit_secrets()

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from crew import legal_assistant_crew

# Page 
st.set_page_config(
    page_title="Federal Eagle 🦅",
    page_icon="⚖️",
    layout="wide",
)

st.title("Federal Eagle 🦅")
st.caption(
    "⚖️ U.S. Federal Legal Assistant powered by CrewAI."
)

# Defaults (no sidebar toggles) 
SHOW_FULL_STATUTE_TEXT = True

# Session State 
if "example" not in st.session_state:
    st.session_state.example = ""
if "analysis_data" not in st.session_state:
    st.session_state.analysis_data = None
if "last_input" not in st.session_state:
    st.session_state.last_input = ""

# Helpers 
def safe_json_loads(raw: str):
    raw = str(raw or "").strip().replace("\r\n", "\n")

    # Strip code fences if any
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.IGNORECASE).strip()
    raw = re.sub(r"\s*```$", "", raw).strip()

    # Normalize smart quotes + trailing commas
    raw = raw.replace("“", '"').replace("”", '"').replace("’", "'")
    raw = re.sub(r",(\s*[}\]])", r"\1", raw)

    try:
        return json.loads(raw)
    except json.JSONDecodeError:
        pass

    # Fallback: try extracting the biggest JSON object/array
    obj_match = re.search(r"\{.*\}", raw, flags=re.DOTALL)
    arr_match = re.search(r"\[.*\]", raw, flags=re.DOTALL)

    candidates = []
    if obj_match:
        candidates.append(obj_match.group(0))
    if arr_match:
        candidates.append(arr_match.group(0))

    candidates.sort(key=len, reverse=True)

    for cand in candidates:
        cand = re.sub(r",(\s*[}\]])", r"\1", cand)
        try:
            return json.loads(cand)
        except json.JSONDecodeError:
            continue

    raise json.JSONDecodeError("Could not parse JSON", raw, 0)


def create_docx_from_text(content: str) -> BytesIO:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.1)
        section.right_margin = Inches(1.1)

    lines = str(content).split("\n")
    for line in lines:
        line = line.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue

        if line.strip().isupper() and len(line.strip()) >= 4:
            p = doc.add_paragraph(line.strip())
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(12)
        else:
            p = doc.add_paragraph(line)
            if p.runs:
                p.runs[0].font.size = Pt(11)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def render_document_preview(text: str):
    lines = str(text).split("\n")
    out = []
    for ln in lines:
        s = ln.strip()
        esc = s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if s.isupper() and len(s) >= 4:
            out.append(f"<div style='font-weight:700; margin-top:10px'>{esc}</div>")
        else:
            out.append(f"<div style='margin:2px 0'>{esc}</div>")

    html = "\n".join(out)

    st.markdown(
        f"""
        <div style="
            border: 1px solid rgba(255,255,255,0.10);
            border-radius: 14px;
            padding: 18px;
            background: rgba(255,255,255,0.03);
            line-height: 1.55;
            font-size: 15px;
        ">{html}</div>
        """,
        unsafe_allow_html=True,
    )


def normalize_filename(name: str) -> str:
    name = (name or "legal_document").strip().lower()
    name = re.sub(r"[^a-z0-9_-]+", "_", name)
    return name.strip("_") or "legal_document"


# Sidebar 
with st.sidebar:
    st.subheader("ℹ️ About this tool")
    st.write(
        "Describe a situation that may involve U.S. federal law. The app summarizes the issue, "
        "retrieves relevant U.S. Code (USC) sections from a local semantic database, optionally pulls "
        "federal precedents from trusted legal sources, and generates a practical draft document for "
        "internal use."
    )

    st.divider()
    st.subheader("📋 Example Cases")

    examples = {
        "🖥️ Computer Fraud": (
            "An employee accessed company databases using stolen credentials and "
            "downloaded proprietary customer data. They transmitted this data via "
            "email to a competitor in another state."
        ),
        "💰 Wire Fraud": (
            "A person created a fake investment website and sent promotional emails "
            "to potential investors across multiple states, promising guaranteed returns."
        ),
        "💳 Identity Theft": (
            "Someone used my stolen Social Security number to open credit card accounts "
            "and made fraudulent purchases across several states."
        ),
        "🚨 Bank Robbery": (
            "Two individuals entered a federally-insured bank with firearms and stole cash. "
            "They fled across state lines and were apprehended later."
        ),
        "🏦 Bank Fraud": (
            "False information was submitted to a federally insured bank to obtain loans "
            "and lines of credit, resulting in financial losses."
        ),
        "💊 Drug Trafficking": (
            "An individual was caught transporting controlled substances across state lines "
            "using interstate highways."
        ),
        "💸 Tax Evasion": (
            "A business owner failed to report income and created false records to "
            "claim fraudulent deductions over several years."
        ),
        "💵 Money Laundering": (
            "Funds from suspected unlawful activity were moved through multiple accounts "
            "and structured transactions to conceal their source."
        ),
        "📨 Mail Fraud": (
            "A person mailed deceptive documents to victims to induce payments for services "
            "that were never provided."
        ),
        "🔐 Internal Data Breach": (
            "A company discovered unauthorized access to internal systems and is conducting an "
            "internal review for compliance, reporting obligations, and containment."
        ),
        "🌱 Environmental Violation": (
            "A company discharged pollutants into a waterway and may have violated federal "
            "environmental regulations."
        ),
        "📑 Federal Contract Fraud": (
            "A contractor submitted inflated invoices and false certifications on a federal "
            "government contract."
        ),
        "🧑‍💼 Workplace Harassment": (
            "An employee at a federal contractor reports harassment and retaliation after making "
            "an internal complaint, and wants to document the facts and options."
        ),
    }

    for label, text in examples.items():
        if st.button(label, use_container_width=True):
            st.session_state.example = text

# Input 
default_text = st.session_state.get("example", "") or st.session_state.get("last_input", "")

# Public-use safety controls (kill switch, mandatory disclaimer, rate limit)
from tools.reliability import is_enabled, disabled_message, SessionRateLimiter
from tools.safety import check_input
from tools.observability import new_run, get_logger, StageTimer

if "session_id" not in st.session_state:
    import uuid
    st.session_state.session_id = uuid.uuid4().hex
if "_rate_limiter" not in st.session_state:
    st.session_state._rate_limiter = SessionRateLimiter()

if not is_enabled():
    st.error(f"🛑 {disabled_message()}")
    st.stop()

# Mandatory disclaimer acknowledgement (a public legal tool needs this)
st.markdown(
    "> ⚠️ **Educational tool only.** Federal Eagle does not provide legal advice. "
    "If you face a real legal matter, please consult a licensed attorney in your jurisdiction."
)
ack = st.checkbox(
    "I understand this is an educational tool, not legal advice, and I will consult a lawyer for real matters.",
    value=False,
    key="disclaimer_ack",
)

with st.form("legal_form"):
    user_input = st.text_area(
        "📝 Describe your federal legal issue:",
        value=default_text,
        height=190,
        placeholder="What happened? Who was involved? Where/when? Any interstate or federal elements?",
    )
    submitted = st.form_submit_button("🔍 Analyze Case", use_container_width=True)

# Run analysis ONLY on submit
if submitted:
    if not ack:
        st.warning("⚠️ Please acknowledge the educational-use disclaimer above before submitting.")
        st.stop()
    if not user_input.strip():
        st.warning("⚠️ Please enter a legal issue to analyze.")
    else:
        # Rate limit per session
        allowed, remaining, retry_after = st.session_state._rate_limiter.check(st.session_state.session_id)
        if not allowed:
            st.error(
                f"🛑 Rate limit reached. You can submit another analysis in about {retry_after} seconds."
            )
            st.stop()

        # Allocate a fresh run_id and logger for this analysis
        run_id = new_run(label=f"streamlit_{st.session_state.session_id[:8]}")
        log = get_logger(run_id)
        log("request.received", payload={
            "session_id": st.session_state.session_id,
            "input_chars": len(user_input),
            "remaining_quota": remaining,
        })

        # Safety gates BEFORE we spend a single token
        decision = check_input(user_input)
        log("safety.decision", payload=decision.as_log_payload())
        if not decision.allowed:
            st.error(
                f"🛑 Federal Eagle cannot analyze that input.\n\n**Reason:** {decision.reason}\n\n"
                "If you are in immediate danger, please contact local authorities. "
                "For mental health support in the U.S., dial 988."
            )
            st.stop()
        clean_input = decision.sanitized_input or user_input
        if decision.category == "injection_redacted":
            st.info("ℹ️ Detected and redacted prompt-injection patterns in your input.")

        st.session_state.example = ""
        st.session_state.last_input = user_input
        st.session_state.last_run_id = run_id

        with st.spinner("Running analysis…"):
            with StageTimer(log, "crew", payload={"input_chars": len(clean_input)}):
                result = legal_assistant_crew.kickoff(inputs={"user_input": clean_input})
            raw_result = str(result)
            log("crew.raw_output", payload={"chars": len(raw_result)})

        try:
            data = safe_json_loads(raw_result)
            # Deterministic excerpt repair: replace any drafter-paraphrased statute
            # excerpts with verbatim substrings of upstream USC content. No LLM cost.
            try:
                from tools.usc_sections_search_tool import repair_drafter_excerpts
                usc_out = getattr(result, "tasks_output", None)
                upstream_top: list = []
                if usc_out and len(usc_out) >= 2:
                    usc_raw = getattr(usc_out[1], "raw", None) or str(usc_out[1])
                    try:
                        upstream_top = (safe_json_loads(usc_raw) or {}).get("top_statutes", []) or []
                    except Exception:
                        upstream_top = []
                if isinstance(data, dict) and upstream_top:
                    data = repair_drafter_excerpts(data, upstream_top)
                    log("drafter.excerpts_repaired", payload={
                        "n_repaired": data.get("__excerpts_repaired__", 0),
                    })
            except Exception as e:
                log("drafter.repair_failed", payload={"error": str(e)}, level="warn")
            st.session_state.analysis_data = data
            log("request.completed", payload={"ok": True})
        except json.JSONDecodeError as je:
            log("error.invalid_json", payload={"error": str(je)}, level="error")
            st.error(f"❌ Model returned invalid JSON: {je}")
            st.code(raw_result, language="text")
            st.session_state.analysis_data = None


# Render from persisted state (survives download reruns) 
data = st.session_state.analysis_data
if data:
    disclaimer = (data.get("disclaimer") or "").strip()
    if disclaimer:
        if not disclaimer.lstrip().startswith("⚠️"):
            disclaimer = f"⚠️ {disclaimer}"
        st.info(disclaimer)

    st.divider()

    tab_names = ["🔍 Summary", "📚 Statutes (USC DB)", "🧩 Elements", "🏛️ Precedents", "📄 Draft"]
    tabs = st.tabs(tab_names)

    # Summary 
    with tabs[0]:
        summary = data.get("summary", {}) or {}
        st.subheader(summary.get("primary_issue", "Case Summary"))

        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("Case Type", summary.get("case_type", "unclear"))
        with col2:
            st.metric("Confidence", summary.get("confidence", "medium"))
        with col3:
            hooks = summary.get("federal_hook", [])
            st.metric("Federal Hooks", str(len(hooks)) if isinstance(hooks, list) else "—")

        hooks = summary.get("federal_hook", [])
        if hooks:
            st.markdown("**Why this may be federal:**")
            for h in hooks:
                st.write(f"- {h}")

        assumptions = summary.get("assumptions", [])
        if assumptions:
            with st.expander("Assumptions used", expanded=False):
                for a in assumptions:
                    st.write(f"- {a}")

        questions = data.get("clarifying_questions", [])
        if questions:
            st.markdown("**Clarifying questions (to improve accuracy):**")
            for q in questions:
                st.write(f"- {q}")

        next_steps = data.get("next_steps", [])
        if next_steps:
            st.markdown("**Practical next steps:**")
            for s in next_steps:
                st.write(f"- {s}")

    # Statutes 
    with tabs[1]:
        statutes = data.get("statutes", []) or []
        if not statutes:
            st.warning("No statutes returned.")
        else:
            rows = []
            for s in statutes:
                rows.append(
                    {
                        "Citation": s.get("citation", ""),
                        "Section": s.get("section_title", ""),
                        "Title": s.get("title", ""),
                        "Why relevant": s.get("why_relevant", ""),
                    }
                )
            st.dataframe(rows, use_container_width=True, hide_index=True)

            for s in statutes:
                citation = s.get("citation", "USC")
                title_name = s.get("title_name", "")
                section_title = s.get("section_title", "")

                label = f"{citation} — {section_title}" if section_title else citation
                if title_name:
                    label += f" ({title_name})"

                with st.expander(label):
                    st.markdown(f"**Why relevant:** {s.get('why_relevant', '')}")

                    elems = s.get("elements", [])
                    if elems:
                        st.markdown("**High-level elements (summary):**")
                        for e in elems:
                            st.write(f"- {e}")

                    excerpt = (s.get("excerpt") or "").strip()
                    if excerpt:
                        st.markdown("**Excerpt (retrieved from USC text):**")
                        st.code(excerpt, language="text")

                    if SHOW_FULL_STATUTE_TEXT:
                        full_text = (s.get("content") or "").strip()
                        if full_text:
                            st.markdown("**Full retrieved statute text:**")
                            st.code(full_text[:6000], language="text")

    # Elements 
    with tabs[2]:
        elements_analysis = data.get("elements_analysis", []) or []
        if not elements_analysis:
            st.warning("No elements analysis returned.")
        else:
            for item in elements_analysis:
                citation = item.get("citation", "Statute")
                with st.expander(f"{citation}", expanded=True):
                    checklist = item.get("checklist", []) or []
                    if not checklist:
                        st.write("No checklist items.")
                        continue

                    for idx, c in enumerate(checklist):
                        status = c.get("status", "unknown")
                        icon = "✅" if status == "met" else ("❓" if status == "unknown" else "⚠️")
                        st.markdown(f"{icon} **{c.get('element', '')}**  \nStatus: `{status}`")

                        facts = c.get("supporting_facts", []) or []
                        for f in facts:
                            st.write(f"- {f}")

                        if idx != len(checklist) - 1:
                            st.divider()

    # Precedents 
    with tabs[3]:
        precedents = data.get("precedents", []) or []
        notes = data.get("precedent_notes", None)
        if notes is None:
            notes = data.get("notes", [])
        if not isinstance(notes, list):
            notes = [str(notes)]

        if not precedents:
            st.warning("No precedents returned.")
            if notes:
                st.caption("Notes:")
                for n in notes:
                    if str(n).strip():
                        st.write(f"- {n}")
        else:
            if notes:
                with st.expander("Notes", expanded=False):
                    for n in notes:
                        if str(n).strip():
                            st.write(f"- {n}")

            for p in precedents:
                name = p.get("name", "Case")
                court_year = p.get("court_year", "")
                header = f"{name} ({court_year})" if court_year else name

                with st.expander(header):
                    holding = (p.get("holding") or "").strip()
                    relevance = (p.get("relevance") or "").strip()
                    url = (p.get("url") or "").strip()

                    if holding:
                        st.markdown(f"**Holding:** {holding}")
                    if relevance:
                        st.markdown(f"**Relevance:** {relevance}")
                    if url:
                        st.link_button("Open source", url, use_container_width=True)

    # Draft 
    with tabs[4]:
        draft = data.get("draft_document", {}) or {}
        doc_type = draft.get("document_type", "Draft Document")
        content = (draft.get("content") or "").strip()

        st.subheader(doc_type)

        if not content:
            st.warning("No drafted document content returned.")
        else:
            render_document_preview(content)
            
            st.markdown("<div style='height:18px'></div>", unsafe_allow_html=True)

            safe_name = normalize_filename(doc_type)

            col1, col2 = st.columns(2)
            with col1:
                bio = create_docx_from_text(content)
                st.download_button(
                    "📄 Download Word (.docx)",
                    data=bio,
                    file_name=f"{safe_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            with col2:
                st.download_button(
                    "📝 Download Text (.txt)",
                    content,
                    file_name=f"{safe_name}.txt",
                    use_container_width=True,
                )
